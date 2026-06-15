"""
RakshAI — Comprehensive Architecture & Client Guide (DOCX Generator)
Generates a professional .docx document covering every system layer,
agent workflow, rule engine pipeline, data flow, and API surface.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
from datetime import datetime


def set_cell_shading(cell, color_hex: str):
    """Set background shading for a table cell."""
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def add_styled_table(doc, headers, rows, col_widths=None):
    """Add a formatted table with header styling."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(cell, "1B2A4A")

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, value in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(value)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
            if r_idx % 2 == 1:
                set_cell_shading(cell, "EDF2F7")

    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)

    doc.add_paragraph()
    return table


def build_document():
    doc = Document()

    # ── Document-level style defaults ───────────────────────────────────────
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = RGBColor(30, 30, 30)

    for level in range(1, 5):
        heading_style = doc.styles[f"Heading {level}"]
        heading_style.font.color.rgb = RGBColor(27, 42, 74)  # Dark navy

    # ════════════════════════════════════════════════════════════════════════
    # COVER PAGE
    # ════════════════════════════════════════════════════════════════════════
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("RakshAI")
    run.font.size = Pt(42)
    run.bold = True
    run.font.color.rgb = RGBColor(27, 42, 74)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Automated Web Application Penetration Testing Platform")
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(80, 80, 80)

    doc.add_paragraph()

    tagline = doc.add_paragraph()
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = tagline.add_run("Complete Architecture & Client Guide")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(100, 100, 100)
    run.italic = True

    for _ in range(4):
        doc.add_paragraph()

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(f"Document Version: 1.0\nDate: {datetime.now().strftime('%B %d, %Y')}\nClassification: CONFIDENTIAL — Client Use Only")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(120, 120, 120)

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # TABLE OF CONTENTS (placeholder)
    # ════════════════════════════════════════════════════════════════════════
    doc.add_heading("Table of Contents", level=1)
    toc_items = [
        "1.  Executive Summary",
        "2.  Platform Overview & Core Principles",
        "3.  System Architecture Overview",
        "4.  Layer 1 — Frontend (React Dashboard)",
        "5.  Layer 2 — Backend API (FastAPI)",
        "6.  Layer 3 — Orchestrator Service",
        "7.  Layer 4 — Multi-Agent Pipeline (LangGraph)",
        "    7.1  Coordinator Agent",
        "    7.2  Reconnaissance Agent",
        "    7.3  Attack Strategy Agent",
        "    7.4  Exploit Execution Agent",
        "    7.5  Validation Agent",
        "    7.6  PoC Generation Agent",
        "    7.7  Remediation Agent",
        "8.  Layer 5 — Deterministic Rule Engine",
        "    8.1  Context Normalizer",
        "    8.2  OWASP Mapper",
        "    8.3  Test Case Evaluator",
        "    8.4  Payload Binder",
        "    8.5  Safety Enforcer",
        "    8.6  Validator Selector",
        "    8.7  Attack Plan Generator",
        "9.  Knowledge Base (YAML Rules Repository)",
        "10. Data Layer & Storage",
        "11. Real-Time Communication (WebSockets)",
        "12. LLM Integration & AI Strategy",
        "13. Security & Safety Mechanisms",
        "14. OWASP Top 10:2025 Coverage Matrix",
        "15. Complete Scan Workflow — Step by Step",
        "16. API Endpoint Reference",
        "17. Technology Stack Summary",
        "18. Deployment Architecture",
        "19. Glossary of Terms",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(0)
        for run in p.runs:
            run.font.size = Pt(10)

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # 1. EXECUTIVE SUMMARY
    # ════════════════════════════════════════════════════════════════════════
    doc.add_heading("1. Executive Summary", level=1)
    doc.add_paragraph(
        "RakshAI is an enterprise-grade Automated Web Application Penetration Testing platform "
        "that delivers continuous, repeatable, and explainable security assessments. Unlike "
        "traditional vulnerability scanners that rely on signature matching, or AI-only tools "
        "that can produce hallucinated results, RakshAI uniquely combines:"
    )
    bullets = [
        "Deterministic, YAML-driven Rule Engine — All security decisions (what to test, which payloads to use, how to validate) are governed by expert-crafted YAML rules. This guarantees that every test is explainable, safe, and 100% repeatable.",
        "Multi-Agent AI Architecture — Seven specialized autonomous agents work together through a state-machine pipeline (powered by LangGraph) to execute the full penetration testing lifecycle automatically.",
        "LLM-Enhanced Reporting — Large Language Models are used exclusively for generating human-readable business impact analyses, remediation guidance, and executive summaries — never for making security detection decisions.",
        "Zero-Hallucination Validation — Every finding is replayed 3 times independently, requiring an 85% reproducibility threshold before being marked as confirmed.",
    ]
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")

    doc.add_paragraph(
        "The platform transforms the traditionally manual, time-intensive penetration testing process "
        "into a fully automated, continuously running security workflow — enabling organizations to "
        "secure their web applications rapidly and with measurable confidence."
    )

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # 2. PLATFORM OVERVIEW & CORE PRINCIPLES
    # ════════════════════════════════════════════════════════════════════════
    doc.add_heading("2. Platform Overview & Core Principles", level=1)

    doc.add_heading("2.1 Core Design Principles", level=2)
    principles = [
        ("Deterministic Security Decisions", "All vulnerability detection, payload selection, and validation logic is driven by structured YAML rules — not by LLM inference. This ensures every security decision is explainable, auditable, and repeatable across scans."),
        ("AI for Explanation, Not Detection", "LLMs (via Ollama, running locally) are used strictly for translating technical findings into business-language reports, generating remediation code snippets, and creating executive summaries. They never decide whether a vulnerability exists."),
        ("Zero-Hallucination Guardrails", "Every raw finding must pass a 3× replay validation with an 85% confidence threshold. Findings that fail are automatically classified as false positives with an LLM-generated explanation of why."),
        ("Privacy-First AI", "All LLM processing runs locally via Ollama on the customer's infrastructure. No vulnerability data, scan results, or target information ever leaves the deployment environment."),
        ("Enterprise Safety Controls", "A dedicated Safety Enforcer component in the Rule Engine blocks destructive payloads (DROP TABLE, rm -rf, etc.) in production environments. Rate limiting, scope validation, and policy enforcement protect target systems from accidental denial of service."),
    ]
    for title, desc in principles:
        p = doc.add_paragraph()
        run = p.add_run(f"{title}: ")
        run.bold = True
        p.add_run(desc)

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # 3. SYSTEM ARCHITECTURE OVERVIEW
    # ════════════════════════════════════════════════════════════════════════
    doc.add_heading("3. System Architecture Overview", level=1)
    doc.add_paragraph(
        "RakshAI is built as a layered, microservice-oriented system where each layer has clearly "
        "defined responsibilities. The following diagram describes the high-level architecture:"
    )

    arch_flow = doc.add_paragraph()
    arch_flow.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = arch_flow.add_run(
        "┌─────────────────────────────────────────────────────────────┐\n"
        "│                   FRONTEND (React + TypeScript)             │\n"
        "│   Dashboard │ Scans │ Findings │ Reports │ Attack Surface   │\n"
        "└──────────────────────────┬──────────────────────────────────┘\n"
        "                           │ REST API + WebSocket\n"
        "┌──────────────────────────▼──────────────────────────────────┐\n"
        "│                   BACKEND API (FastAPI)                      │\n"
        "│   /api/v1/scans │ /vulnerabilities │ /reports │ /dashboard   │\n"
        "└──────────────────────────┬──────────────────────────────────┘\n"
        "                           │\n"
        "┌──────────────────────────▼──────────────────────────────────┐\n"
        "│                ORCHESTRATOR SERVICE (Layer 2)                │\n"
        "│   Scope Validation → Policy → Rate Limiting → Launch        │\n"
        "└──────────────────────────┬──────────────────────────────────┘\n"
        "                           │\n"
        "┌──────────────────────────▼──────────────────────────────────┐\n"
        "│         COORDINATOR AGENT (LangGraph State Machine)         │\n"
        "│  Planning → Recon → Strategy → Executor → Validator → PoC  │\n"
        "└──────────────────────────┬──────────────────────────────────┘\n"
        "                           │\n"
        "┌──────────────────────────▼──────────────────────────────────┐\n"
        "│              RULE ENGINE (7 Components, YAML)               │\n"
        "│  Normalizer → OWASP Mapper → Test Cases → Payloads →       │\n"
        "│  Safety Enforcer → Validator Selector → Attack Plan         │\n"
        "└──────────────────────────┬──────────────────────────────────┘\n"
        "                           │\n"
        "┌──────────────────────────▼──────────────────────────────────┐\n"
        "│                     DATA LAYER                               │\n"
        "│   PostgreSQL/SQLite │ Redis │ MinIO │ Ollama (Local LLM)    │\n"
        "└─────────────────────────────────────────────────────────────┘"
    )
    run.font.size = Pt(8)
    run.font.name = "Consolas"

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # 4. LAYER 1 — FRONTEND
    # ════════════════════════════════════════════════════════════════════════
    doc.add_heading("4. Layer 1 — Frontend (React Dashboard)", level=1)
    doc.add_paragraph(
        "The frontend is a modern, single-page application built with React 18, TypeScript 5, "
        "and Tailwind CSS with shadcn/ui components. It provides the complete user interface for "
        "managing scans, viewing results, and generating reports."
    )

    doc.add_heading("4.1 Frontend Pages & Features", level=2)
    add_styled_table(doc,
        ["Page", "Purpose", "Key Features"],
        [
            ["Dashboard", "Real-time overview of security posture", "Active scan count, severity breakdown charts, recent findings, trend analysis"],
            ["Scans", "Create and manage penetration tests", "New scan wizard, scan list with filters, start/stop controls, progress tracking"],
            ["Scan Detail", "Deep dive into a specific scan", "Live phase tracker, endpoint list, finding count by severity, log viewer"],
            ["Vulnerabilities", "Browse all discovered findings", "Severity filters, OWASP category grouping, false positive management"],
            ["Vulnerability Detail", "Full finding analysis", "Description, evidence, PoC steps with screenshots, cURL commands, remediation guide"],
            ["Reports", "Generate and download reports", "PDF, Word, Excel export, executive summary, detailed technical appendix"],
            ["Attack Surface", "Visual map of discovered endpoints", "Endpoint tree, method breakdown, parameter analysis, auth indicators"],
            ["Evidence Viewer", "Browse scan evidence files", "Screenshots, HTTP traces, HAR files, raw tool outputs"],
            ["Governance", "Policy and compliance management", "Scan policies, schedule configuration, compliance mapping"],
            ["Audit Trail", "Complete activity log", "Who did what and when, all scan actions with timestamps"],
            ["Diagnostics", "System health monitoring", "Backend status, Ollama connectivity, database health, log viewer"],
        ],
        col_widths=[3.5, 4.5, 8.0],
    )

    doc.add_heading("4.2 Real-Time Updates", level=2)
    doc.add_paragraph(
        "The frontend maintains a persistent WebSocket connection (/ws/{client_id}) to the backend. "
        "As each agent in the pipeline completes a phase, progress events are pushed in real-time "
        "to the dashboard — the user sees live phase transitions (\"Discovering → Strategizing → "
        "Exploiting → Validating → Generating PoC → Completed\") without page refresh."
    )

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # 5. LAYER 2 — BACKEND API
    # ════════════════════════════════════════════════════════════════════════
    doc.add_heading("5. Layer 2 — Backend API (FastAPI)", level=1)
    doc.add_paragraph(
        "The backend is built with FastAPI, providing a high-performance asynchronous REST API. "
        "It handles all client requests, manages the database, and serves as the gateway to the "
        "orchestration and agent layers."
    )

    doc.add_heading("5.1 Key Backend Components", level=2)
    add_styled_table(doc,
        ["Component", "File", "Responsibility"],
        [
            ["Application Entry", "app/main.py", "FastAPI app initialization, middleware (CORS, GZip), Prometheus metrics, WebSocket endpoint"],
            ["API Router", "app/api/v1/", "Versioned REST endpoints organized by domain (scans, vulnerabilities, reports, etc.)"],
            ["Database Models", "app/models/models.py", "SQLAlchemy ORM models: Scan, Endpoint, Vulnerability, Report, AuditLog, ScheduledScan"],
            ["Pydantic Schemas", "app/models/schemas.py", "Request/response validation schemas with Pydantic v2"],
            ["Configuration", "app/core/config.py", "Environment-based settings (DB URL, Redis, Ollama, CORS, etc.)"],
            ["Database Engine", "app/core/database.py", "SQLAlchemy engine, session factory, SQLite/PostgreSQL support"],
            ["WebSocket Manager", "app/core/websocket_manager.py", "Client connection tracking and broadcast for real-time events"],
            ["Redis Client", "app/core/redis_client.py", "Pub/Sub listener for inter-service communication"],
            ["Knowledge Base", "app/core/knowledge_base.py", "YAML rule loader for the rule engine"],
        ],
        col_widths=[3.5, 4.0, 8.5],
    )

    doc.add_heading("5.2 Startup Lifecycle", level=2)
    startup_steps = [
        "FastAPI application is created with CORS, GZip, and Prometheus middleware.",
        "On startup event: database tables are created (create_all), SQLite schema is ensured.",
        "Redis Pub/Sub listener starts for inter-service communication.",
        "Scan scheduler service starts for recurring/scheduled scans.",
        "The API router is mounted at /api/v1/ with all endpoint modules.",
        "Prometheus metrics endpoint is mounted at /metrics for monitoring.",
    ]
    for i, step in enumerate(startup_steps, 1):
        doc.add_paragraph(f"{i}. {step}")

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # 6. LAYER 3 — ORCHESTRATOR SERVICE
    # ════════════════════════════════════════════════════════════════════════
    doc.add_heading("6. Layer 3 — Orchestrator Service", level=1)
    doc.add_paragraph(
        "The Orchestrator is the enterprise control plane sitting between the API layer and the "
        "agent pipeline. It enforces critical safeguards before any scan begins."
    )

    doc.add_heading("6.1 Pre-Scan Validation Pipeline", level=2)
    add_styled_table(doc,
        ["Phase", "What It Does", "Failure Behavior"],
        [
            ["1. Scope Validation", "Checks if the target URL is authorized. Blocks scanning of blacklisted domains (google.com, facebook.com, etc.) and localhost.", "Returns SCOPE_VIOLATION error; scan is not started."],
            ["2. Policy Enforcement", "Loads the effective scan policy (default + user overrides). Checks time windows, forbidden days, and production-environment restrictions.", "Returns POLICY_VIOLATION error; scan is not started."],
            ["3. Rate Limiting", "Enforces maximum 5 scans per hour per target and minimum 2-minute gap between scans of the same target.", "Returns RATE_LIMIT_EXCEEDED error; user must wait."],
            ["4. Resource Check", "Checks if there are available slots (max 5 concurrent scans). If not, the scan is added to a FIFO queue.", "Scan is queued with a position number; auto-launched when a slot opens."],
            ["5. Scan Launch", "Creates the CoordinatorAgent, registers the scan as active, and launches it as an asyncio background task.", "On failure, scan is marked FAILED in the database with an error message."],
        ],
        col_widths=[3.0, 6.0, 7.0],
    )

    doc.add_heading("6.2 Policy Configuration", level=2)
    doc.add_paragraph("Each scan operates under an effective policy that controls:")
    policy_items = [
        "max_depth: How many levels deep crawling should go (0 = single page, 3 = full site)",
        "max_endpoints: Maximum number of endpoints to discover (1–100)",
        "max_attacks: Maximum number of attack tests to execute (up to 50)",
        "allowed_attacks: Which attack types are permitted (IDOR, XSS, SQLI, AUTH_BYPASS)",
        "forbidden_attacks: Which attacks are always blocked (DOS, RESOURCE_EXHAUSTION)",
        "rate_limit: Requests per second and concurrent request limits",
        "time_window: Allowed hours and forbidden days of the week",
    ]
    for item in policy_items:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # 7. LAYER 4 — MULTI-AGENT PIPELINE
    # ════════════════════════════════════════════════════════════════════════
    doc.add_heading("7. Layer 4 — Multi-Agent Pipeline (LangGraph)", level=1)
    doc.add_paragraph(
        "The heart of RakshAI is a multi-agent system powered by LangGraph — a library for "
        "building stateful, graph-based AI agent workflows. Seven specialized agents execute "
        "in a defined sequence, sharing state through a typed dictionary (PentestState) and "
        "persisting results to PostgreSQL at each step."
    )

    flow = doc.add_paragraph()
    flow.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = flow.add_run(
        "[Planning] → [Recon] → [Strategy] → [Executor] → [Validator] → [PoC] → END"
    )
    run.bold = True
    run.font.size = Pt(12)

    doc.add_paragraph()

    # 7.1 Coordinator
    doc.add_heading("7.1 Coordinator Agent", level=2)
    doc.add_paragraph("File: app/agents/coordinator.py")
    doc.add_paragraph(
        "The Coordinator is the \"brain\" that builds and executes the LangGraph state machine. "
        "It does not perform scanning itself — instead, it orchestrates the six downstream agents."
    )
    coord_responsibilities = [
        "Loads the scan record from PostgreSQL and resolves the scan UUID.",
        "Creates the LangGraph StateGraph with 6 nodes and linear edges.",
        "Seeds the initial PentestState with scan metadata.",
        "Invokes the graph asynchronously (graph.ainvoke).",
        "Updates scan status and phase in the database after each node completes.",
        "Handles errors gracefully — if any agent fails, the scan is marked FAILED.",
        "Falls back to sequential execution if LangGraph is not installed.",
    ]
    for r in coord_responsibilities:
        doc.add_paragraph(r, style="List Bullet")

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("Shared State Schema (PentestState):")
    run.bold = True
    add_styled_table(doc,
        ["Key", "Type", "Written By", "Purpose"],
        [
            ["scan_id", "str", "Planning", "UUID identifying the scan across all agents"],
            ["target_url", "str", "Planning", "The target URL being scanned"],
            ["policy", "Dict", "Planning", "Effective scan policy (limits, allowed attacks)"],
            ["strategy", "Dict", "Planning", "LLM-generated strategic analysis"],
            ["recon", "Dict", "Recon", "Discovered endpoints, technologies, forms"],
            ["strategy_plan", "Dict", "Strategy", "Prioritized attack vectors"],
            ["execution", "Dict", "Executor", "Raw findings from exploit execution"],
            ["validation", "Dict", "Validator", "Validation results with confidence scores"],
            ["poc", "Dict", "PoC", "Generated proof-of-concept evidence"],
        ],
        col_widths=[3.0, 2.0, 2.5, 8.5],
    )

    # 7.2 Recon
    doc.add_heading("7.2 Reconnaissance Agent", level=2)
    doc.add_paragraph("File: app/agents/recon.py")
    doc.add_paragraph(
        "The Recon Agent discovers the target's attack surface by crawling, probing, and fingerprinting."
    )
    add_styled_table(doc,
        ["Phase", "Tool Used", "What It Discovers"],
        [
            ["1. HTTP Probing", "httpx", "Status codes, page titles, web server type, TLS info"],
            ["2. Web Crawling", "Katana", "All reachable URLs, JavaScript routes, API endpoints"],
            ["3. Technology Detection", "httpx + custom", "Frameworks (React, Django, Laravel), server software"],
            ["4. Template Scanning", "Nuclei", "Known CVEs, exposed credentials, misconfigurations"],
            ["5. Form Discovery", "Katana + Playwright", "HTML forms, input fields, hidden parameters"],
            ["6. Data Persistence", "PostgreSQL", "All endpoints stored as Endpoint records linked to the Scan"],
            ["7. Raw Output Upload", "MinIO", "Full tool outputs stored for audit trail"],
        ],
        col_widths=[3.5, 3.0, 9.5],
    )

    # 7.3 Strategy
    doc.add_heading("7.3 Attack Strategy Agent", level=2)
    doc.add_paragraph("File: app/agents/strategy.py")
    doc.add_paragraph(
        "This agent uses LLM-powered threat modeling to analyze discovered endpoints and create "
        "a prioritized attack plan based on OWASP Top 10 categories."
    )
    strategy_steps = [
        "Loads all discovered endpoints from the database.",
        "Sends an endpoint summary to the LLM with a structured prompt requesting JSON threat model output.",
        "LLM identifies injection points, access control targets, auth bypass candidates, and sensitive endpoints.",
        "Fallback: if LLM is unavailable, a deterministic rule-based threat model is generated using URL pattern matching.",
        "Attack vectors are prioritized: AUTH_BYPASS (priority 95) → IDOR (85) → SQLI (80) → XSS (70).",
        "Attack nodes are created in the database for the Executor to consume.",
    ]
    for i, s in enumerate(strategy_steps, 1):
        doc.add_paragraph(f"{i}. {s}")

    # 7.4 Executor
    doc.add_heading("7.4 Exploit Execution Agent", level=2)
    doc.add_paragraph("File: app/agents/executor.py")
    doc.add_paragraph(
        "The Executor runs the actual security tests against the target, sending payloads and "
        "analyzing responses through a sandboxed tool execution environment."
    )
    add_styled_table(doc,
        ["Attack Type", "Tool", "Detection Method", "Severity"],
        [
            ["SQL Injection (SQLI)", "sqlmap", "Checks for 'is vulnerable' or 'injectable' in output", "HIGH"],
            ["Cross-Site Scripting (XSS)", "dalfox", "Parses JSON output for reflected payloads", "MEDIUM"],
            ["IDOR", "Custom idor_tester", "Checks for 'IDOR_DETECTED' marker", "HIGH"],
            ["Auth Bypass", "Custom auth_bypass_tester", "Checks for 'BYPASS_SUCCESSFUL' marker", "CRITICAL"],
            ["Security Headers", "Custom security_headers", "Checks for missing CSP, HSTS, X-Frame headers", "LOW–MEDIUM"],
        ],
        col_widths=[3.5, 3.5, 5.5, 3.5],
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("Smart Attack Capping: ")
    run.bold = True
    p.add_run(
        "To prevent excessive HTTP traffic, the Executor enforces per-type caps: "
        "max 5 SQLI tests, max 5 XSS tests, max 3 IDOR tests, and max 2 Auth Bypass tests per scan. "
        "Each finding is persisted to the Vulnerability table with OWASP category, CWE ID, and severity."
    )

    # 7.5 Validator
    doc.add_heading("7.5 Validation Agent", level=2)
    doc.add_paragraph("File: app/agents/validator.py")
    doc.add_paragraph(
        "The Validation Agent implements the Zero-Hallucination Guardrail — the core mechanism "
        "that prevents false positives from reaching the final report."
    )
    doc.add_heading("How Validation Works:", level=3)
    validation_steps = [
        "For each unvalidated finding, the agent replays the exact attack 3 times independently.",
        "Each replay re-executes the tool (sqlmap, dalfox, etc.) with the same parameters.",
        "A 2-second delay is enforced between replays to avoid rate limiting interference.",
        "The confidence score = (successes / 3). A score ≥ 0.85 (i.e., 2 out of 3 successes) marks the finding as VALIDATED.",
        "Findings below the threshold are marked FALSE_POSITIVE.",
        "For false positives, the LLM analyzes the replay results and generates an explanation (e.g., 'Tool misdetection due to WAF response caching').",
        "The finding's status, confidence score, and replay count are updated in the database.",
    ]
    for i, s in enumerate(validation_steps, 1):
        doc.add_paragraph(f"{i}. {s}")

    # 7.6 PoC
    doc.add_heading("7.6 PoC Generation Agent", level=2)
    doc.add_paragraph("File: app/agents/poc_generator.py")
    doc.add_paragraph(
        "For every validated finding, this agent generates comprehensive proof-of-concept evidence "
        "that enables developers to understand, reproduce, and fix the vulnerability."
    )
    add_styled_table(doc,
        ["PoC Component", "How It's Generated", "Storage"],
        [
            ["Screenshot", "Playwright browser automation (headless Chrome)", "MinIO object storage with pre-signed URL"],
            ["Step-by-Step PoC", "LLM decomposes the vulnerability into 3–5 actionable reproduction steps, each with its own screenshot", "JSON in poc_steps column"],
            ["HTTP Trace", "Raw request/response capture (method, headers, body, status code)", "MinIO (JSON)"],
            ["cURL Command", "Auto-generated reproduction command from finding parameters", "Stored in finding record"],
            ["Business Impact", "LLM-generated executive-level analysis covering: executive summary, attacker capabilities, data at risk, financial impact, compliance implications", "llm_business_impact column"],
            ["Remediation Guide", "LLM-generated step-by-step fix including: root cause, hotfix, permanent code patch (before/after), hardening, verification steps, references", "llm_remediation column"],
        ],
        col_widths=[3.0, 7.0, 6.0],
    )

    # 7.7 Remediation
    doc.add_heading("7.7 Remediation Agent", level=2)
    doc.add_paragraph("File: app/agents/remediation_agent.py")
    doc.add_paragraph(
        "An on-demand agent that generates instant, technology-specific remediation guidance for "
        "any vulnerability. It detects the target technology (PHP, Python, Node.js, Java, ASP.NET) "
        "from endpoint URLs and response evidence, then generates tailored code examples."
    )

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # 8. LAYER 5 — RULE ENGINE
    # ════════════════════════════════════════════════════════════════════════
    doc.add_heading("8. Layer 5 — Deterministic Rule Engine", level=1)
    doc.add_paragraph(
        "The Rule Engine is RakshAI's \"brain\" for security decisions. It is a 7-component pipeline "
        "that transforms raw discovery data into a safe, prioritized, and validated attack plan — "
        "all without any LLM involvement. Every decision is traceable back to a specific YAML rule."
    )

    components = [
        ("8.1 Context Normalizer", "context_normalizer.py",
         "Takes raw discovery output (different formats from httpx, Katana, Nuclei) and transforms it into a standardized internal format that all downstream components can consume."),
        ("8.2 OWASP Mapper", "owasp_mapper.py",
         "Maps each discovered endpoint to relevant OWASP Top 10:2025 categories based on URL patterns, parameter names, and technology fingerprints. For example, an endpoint with 'id' parameter gets mapped to A01 (Broken Access Control)."),
        ("8.3 Test Case Evaluator", "test_case_evaluator.py",
         "Selects the applicable test cases from the Knowledge Base's 96+ YAML files based on the OWASP mapping. Each test case defines what to test, how to test it, and what success/failure looks like."),
        ("8.4 Payload Binder", "payload_binder.py",
         "Binds safe payloads from ~1,000 YAML-defined payload entries to the selected test cases. Payloads are organized by category: injection (SQL, XSS, command), access-control, authentication, cryptographic, SSRF, file-inclusion, and deserialization."),
        ("8.5 Safety Enforcer", "safety_enforcer.py",
         "The critical safety gate. Blocks destructive payloads (DROP TABLE, rm -rf, shutdown, format, etc.) in production environments. Validates that all payloads are within the scan policy's allowed attack types."),
        ("8.6 Validator Selector", "validator_selector.py",
         "Chooses the appropriate detection validators from 6 YAML validator definitions. Each validator knows how to interpret tool output and determine if a vulnerability was successfully exploited (e.g., checking for SQL error patterns, reflected XSS markers)."),
        ("8.7 Attack Plan Generator", "attack_plan_generator.py",
         "Assembles the final attack execution plan by combining the normalized context, OWASP mappings, selected test cases, bound payloads, safety-checked constraints, and validator assignments into a structured execution document."),
    ]
    for title, filename, desc in components:
        doc.add_heading(title, level=2)
        doc.add_paragraph(f"File: app/rule_engine/{filename}")
        doc.add_paragraph(desc)

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # 9. KNOWLEDGE BASE
    # ════════════════════════════════════════════════════════════════════════
    doc.add_heading("9. Knowledge Base (YAML Rules Repository)", level=1)
    doc.add_paragraph(
        "The Knowledge Base is the repository of all security testing rules, payloads, test cases, "
        "and validators. It is stored as structured YAML files, making it easy to audit, version "
        "control, and extend."
    )

    add_styled_table(doc,
        ["Directory", "Contents", "Count"],
        [
            ["test-cases/A01-broken-access-control/", "IDOR, forced browsing, path traversal test definitions", "10+ files"],
            ["test-cases/A02-security-misconfiguration/", "Default configs, verbose errors, directory listing tests", "8+ files"],
            ["test-cases/A03-software-supply-chain/", "Dependency scanning, known vulnerable components", "5+ files"],
            ["test-cases/A04-crypto-failures/", "Weak TLS, exposed secrets, insecure algorithms", "7+ files"],
            ["test-cases/A05-injection/", "SQL, NoSQL, XSS, XXE, Command Injection tests", "15+ files"],
            ["test-cases/A06-insecure-design/", "Business logic flaws, race conditions", "6+ files"],
            ["test-cases/A07-auth-failures/", "Brute force, session management, credential stuffing", "9+ files"],
            ["test-cases/A08-software-data-integrity/", "Unsigned updates, CI/CD pipeline attacks", "5+ files"],
            ["test-cases/A09-logging-alerting/", "Missing logs, insufficient monitoring", "4+ files"],
            ["test-cases/A10-ssrf/", "Server-Side Request Forgery, internal service access", "6+ files"],
            ["payloads/injection/", "SQL injection, XSS, command injection payloads", "100s of entries"],
            ["payloads/access-control/", "IDOR, privilege escalation payloads", "50+ entries"],
            ["payloads/auth/", "Authentication bypass payloads", "40+ entries"],
            ["payloads/ssrf/", "SSRF payloads targeting internal services", "30+ entries"],
            ["validators/", "Response analysis rules (SQL errors, XSS markers, SSRF indicators)", "6 files"],
            ["workflows/", "Multi-step attack workflows (e.g., post-exploitation)", "3+ files"],
            ["reconnaissance/", "Discovery and fingerprinting rules", "5+ files"],
        ],
        col_widths=[5.5, 6.5, 4.0],
    )

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # 10. DATA LAYER
    # ════════════════════════════════════════════════════════════════════════
    doc.add_heading("10. Data Layer & Storage", level=1)

    doc.add_heading("10.1 Database Schema (PostgreSQL / SQLite)", level=2)
    add_styled_table(doc,
        ["Table", "Key Columns", "Relationships"],
        [
            ["Scan", "id, scan_id (UUID), target_url, scan_type, status, progress_percentage, current_phase, strategy (JSON), total_findings, critical/high/medium/low/info counts, started_at, completed_at", "Has many: Endpoints, Vulnerabilities, Reports"],
            ["Endpoint", "id, scan_id (FK), url, method, endpoint_type, parameters (JSON), discovery_method, requires_auth", "Belongs to: Scan; Has many: Vulnerabilities"],
            ["Vulnerability", "id, scan_id (FK), endpoint_id (FK), title, description, severity (ENUM), confidence, owasp_category, vulnerability_type, cwe_id, status, request_payload, response_evidence, remediation, llm_explanation, llm_business_impact, poc_steps (JSON), is_false_positive", "Belongs to: Scan, Endpoint"],
            ["Report", "id, scan_id (FK), report_id (UUID), report_type, file_path, content (TEXT)", "Belongs to: Scan"],
            ["AuditLog", "id, action, actor, target, details (JSON), timestamp", "—"],
            ["ScheduledScan", "id, target_url, schedule (cron), scan_type, enabled", "—"],
        ],
        col_widths=[2.5, 8.0, 5.5],
    )

    doc.add_heading("10.2 Object Storage (MinIO)", level=2)
    doc.add_paragraph("RakshAI uses MinIO-compatible object storage for large binary artifacts:")
    minio_items = [
        "rakshaidb-raw/ — Raw tool outputs (httpx JSON, Katana crawl data, Nuclei findings)",
        "rakshaidb-screenshots/ — Playwright screenshots for each finding and PoC step",
        "rakshaidb-traces/ — HTTP request/response trace recordings",
        "rakshaidb-reports/ — Generated PDF, Word, and Excel reports",
    ]
    for item in minio_items:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("10.3 Redis", level=2)
    doc.add_paragraph(
        "Redis serves as the Pub/Sub message bus for real-time event delivery. When an agent "
        "completes a phase, it publishes a progress event to Redis, which the backend's listener "
        "picks up and broadcasts to connected WebSocket clients."
    )

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # 11. REAL-TIME COMMUNICATION
    # ════════════════════════════════════════════════════════════════════════
    doc.add_heading("11. Real-Time Communication (WebSockets)", level=1)
    doc.add_paragraph(
        "RakshAI provides live scan progress updates via WebSockets. The flow is:"
    )
    ws_steps = [
        "Frontend connects to ws://backend:8000/ws/{client_id} on page load.",
        "During a scan, each agent calls emit_progress() after completing a sub-phase.",
        "The progress event is published to Redis Pub/Sub.",
        "The backend's Redis listener receives the event and broadcasts it via the WebSocket manager.",
        "The frontend receives the event and updates the UI (phase indicator, progress bar, finding count).",
    ]
    for i, s in enumerate(ws_steps, 1):
        doc.add_paragraph(f"{i}. {s}")

    # ════════════════════════════════════════════════════════════════════════
    # 12. LLM INTEGRATION
    # ════════════════════════════════════════════════════════════════════════
    doc.add_heading("12. LLM Integration & AI Strategy", level=1)
    doc.add_paragraph(
        "RakshAI uses local LLMs via Ollama to ensure all data stays within the customer's infrastructure."
    )

    add_styled_table(doc,
        ["LLM Use Case", "Where Used", "Model", "Fallback If Unavailable"],
        [
            ["Strategic Analysis", "Coordinator Agent", "llama3.2:1b (configurable)", "Deterministic strategy based on URL patterns"],
            ["Threat Modeling", "Strategy Agent", "Same", "Rule-based threat model from URL/param pattern matching"],
            ["False Positive Analysis", "Validation Agent", "Same", "Simple log message (no explanation)"],
            ["Business Impact", "PoC Agent", "Same", "Pre-written impact templates per vulnerability type"],
            ["Remediation Guide", "PoC Agent", "Same", "Pre-written code fix templates per vulnerability type"],
            ["PoC Steps", "PoC Agent", "Same", "Generic 5-step reproduction template"],
            ["Instant Remediation", "Remediation Agent", "Same", "Technology-specific templates from knowledge base"],
        ],
        col_widths=[3.0, 3.0, 3.5, 6.5],
    )

    p = doc.add_paragraph()
    run = p.add_run("Key Guarantee: ")
    run.bold = True
    p.add_run(
        "LLMs are NEVER used for vulnerability detection decisions. They only generate human-readable "
        "text AFTER a vulnerability has been deterministically confirmed by the Rule Engine and Validation Agent."
    )

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # 13. SECURITY & SAFETY
    # ════════════════════════════════════════════════════════════════════════
    doc.add_heading("13. Security & Safety Mechanisms", level=1)
    add_styled_table(doc,
        ["Mechanism", "Description"],
        [
            ["Scope Validation", "Prevents scanning unauthorized targets (blacklisted domains, localhost)"],
            ["Policy Enforcement", "Time-window restrictions, production-environment blocking, attack-type whitelisting"],
            ["Rate Limiting", "Max 5 scans/hour per target, min 2-minute gap between scans"],
            ["Safety Enforcer (Rule Engine)", "Blocks destructive payloads (DROP, DELETE, rm, shutdown) in production"],
            ["Concurrent Scan Limits", "Max 5 simultaneous scans to prevent resource exhaustion"],
            ["Tool Sandbox", "All security tools execute in a controlled sandbox (app/core/tool_sandbox.py) with timeout enforcement"],
            ["Audit Trail", "Every action is logged with actor, timestamp, and details"],
            ["Zero-Hallucination Validation", "3× replay requirement with 85% confidence threshold"],
            ["Local LLM Only", "No cloud API calls — all AI processing runs locally via Ollama"],
        ],
        col_widths=[4.0, 12.0],
    )

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # 14. OWASP COVERAGE MATRIX
    # ════════════════════════════════════════════════════════════════════════
    doc.add_heading("14. OWASP Top 10:2025 Coverage Matrix", level=1)
    add_styled_table(doc,
        ["OWASP ID", "Category", "Test Coverage", "Attack Types"],
        [
            ["A01", "Broken Access Control", "✅ Full", "IDOR, forced browsing, path traversal, privilege escalation"],
            ["A02", "Security Misconfiguration", "✅ Full", "Default configs, verbose errors, missing headers, directory listing"],
            ["A03", "Software Supply Chain", "✅ Full", "Dependency scanning, known CVEs, outdated components"],
            ["A04", "Cryptographic Failures", "✅ Full", "Weak TLS, exposed secrets, insecure algorithms, missing HSTS"],
            ["A05", "Injection", "✅ Full", "SQL injection, NoSQL injection, XSS, XXE, command injection"],
            ["A06", "Insecure Design", "✅ Full", "Business logic flaws, race conditions, missing rate limiting"],
            ["A07", "Authentication Failures", "✅ Full", "Auth bypass, brute force, session fixation, weak passwords"],
            ["A08", "Software & Data Integrity", "✅ Partial", "Unsigned updates, CI/CD pipeline vulnerabilities"],
            ["A09", "Logging & Alerting Failures", "✅ Partial", "Missing security logs, insufficient monitoring"],
            ["A10", "Server-Side Request Forgery", "✅ Full", "SSRF to internal services, cloud metadata access"],
        ],
        col_widths=[2.0, 4.0, 2.0, 8.0],
    )

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # 15. COMPLETE SCAN WORKFLOW
    # ════════════════════════════════════════════════════════════════════════
    doc.add_heading("15. Complete Scan Workflow — Step by Step", level=1)
    doc.add_paragraph(
        "This section describes exactly what happens from the moment a user clicks \"Start Scan\" "
        "to the final report delivery."
    )

    workflow_steps = [
        ("User initiates scan", "The user enters a target URL (e.g., https://example.com) in the frontend dashboard and clicks 'Start Scan'. The frontend sends a POST request to /api/v1/scans/ with the target URL and scan type."),
        ("API creates scan record", "The backend creates a new Scan record in the database with status=PENDING, generates a UUID, and returns it to the frontend."),
        ("Orchestrator validates", "The Orchestrator runs 4 pre-flight checks: scope validation (is the target allowed?), policy compliance (is it the right time/day?), rate limiting (too many recent scans?), and resource availability (are there free scan slots?)."),
        ("Coordinator launches", "A CoordinatorAgent is created and launched as an asyncio background task. The LangGraph state machine is built with 6 nodes."),
        ("Planning phase", "The Coordinator loads the scan from DB, generates an LLM-powered attack strategy (app type, priority OWASP categories, estimated endpoints, risk level), and persists it to the scan's strategy JSON column."),
        ("Reconnaissance", "The ReconAgent crawls the target using httpx (HTTP probing), Katana (web crawling), Nuclei (template scanning), and Playwright (form discovery). All discovered endpoints are persisted to the Endpoint table."),
        ("Attack Strategy", "The AttackStrategyAgent sends the discovered endpoints to the LLM for threat modeling. It generates a prioritized list of attacks (AUTH_BYPASS → IDOR → SQLI → XSS) and creates attack nodes in the database."),
        ("Exploit Execution", "The ExploitExecutionAgent builds attacks from endpoints (capped per type), executes them through the tool sandbox (sqlmap, dalfox, idor_tester, auth_bypass_tester, security_headers), and persists raw findings to the Vulnerability table."),
        ("Validation (3× replay)", "The ValidationAgent retrieves all UNVALIDATED findings and replays each one 3 times. Findings with ≥85% confidence are marked VALIDATED; others are marked FALSE_POSITIVE with an LLM explanation."),
        ("PoC Generation", "The PoCAgent generates screenshots (Playwright), step-by-step reproduction instructions (LLM), HTTP traces, cURL commands, business impact analysis (LLM), and detailed remediation guidance (LLM) for every validated finding."),
        ("Scan completion", "The Coordinator updates severity counts from actual DB data, marks the scan as COMPLETED with progress=100%, and records the completion timestamp."),
        ("Report generation", "The user can request PDF, Word, or Excel reports via /api/v1/scans/{id}/generate-report. Reports include executive summary, severity breakdown, detailed findings with evidence, and remediation timelines."),
    ]
    for i, (title, desc) in enumerate(workflow_steps, 1):
        p = doc.add_paragraph()
        run = p.add_run(f"Step {i}: {title}")
        run.bold = True
        doc.add_paragraph(desc)

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # 16. API REFERENCE
    # ════════════════════════════════════════════════════════════════════════
    doc.add_heading("16. API Endpoint Reference", level=1)
    add_styled_table(doc,
        ["Method", "Endpoint", "Description"],
        [
            ["POST", "/api/v1/scans/", "Create and start a new scan"],
            ["GET", "/api/v1/scans/", "List all scans (with filters)"],
            ["GET", "/api/v1/scans/{id}", "Get scan details"],
            ["PATCH", "/api/v1/scans/{id}", "Update scan status/progress"],
            ["POST", "/api/v1/scans/{id}/start", "Start or restart a scan"],
            ["POST", "/api/v1/scans/{id}/stop", "Stop a running scan"],
            ["DELETE", "/api/v1/scans/{id}", "Delete a scan and all data"],
            ["GET", "/api/v1/scans/{id}/endpoints", "Get discovered endpoints"],
            ["GET", "/api/v1/scans/{id}/phase-summary", "Get phase summary"],
            ["GET", "/api/v1/scans/{id}/logs", "Get scan backend logs"],
            ["GET", "/api/v1/scans/{id}/reports", "Get scan reports"],
            ["POST", "/api/v1/scans/{id}/generate-report", "Generate PDF/Word/Excel report"],
            ["GET", "/api/v1/vulnerabilities/", "List all vulnerabilities"],
            ["GET", "/api/v1/vulnerabilities/{id}", "Get vulnerability detail"],
            ["POST", "/api/v1/vulnerabilities/{id}/remediate", "Generate instant remediation"],
            ["GET", "/api/v1/dashboard/stats", "Dashboard statistics"],
            ["GET", "/api/v1/reports/", "List all reports"],
            ["GET", "/api/v1/evidence/{scan_id}", "Get scan evidence files"],
            ["GET", "/api/v1/audit/", "Get audit trail logs"],
            ["GET", "/api/v1/governance/policies", "Get scan policies"],
            ["GET", "/api/v1/diagnostics/health", "System health check"],
            ["WS", "/ws/{client_id}", "WebSocket for real-time updates"],
            ["GET", "/health", "Backend health check"],
            ["GET", "/metrics", "Prometheus metrics"],
        ],
        col_widths=[2.0, 6.0, 8.0],
    )

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # 17. TECHNOLOGY STACK
    # ════════════════════════════════════════════════════════════════════════
    doc.add_heading("17. Technology Stack Summary", level=1)
    add_styled_table(doc,
        ["Layer", "Technology", "Version", "Purpose"],
        [
            ["Frontend", "React", "18+", "UI component framework"],
            ["Frontend", "TypeScript", "5+", "Type-safe JavaScript"],
            ["Frontend", "Tailwind CSS + shadcn/ui", "Latest", "Styling and UI components"],
            ["Frontend", "TanStack Query v5", "Latest", "Server state management"],
            ["Backend", "Python", "3.11+", "Primary backend language"],
            ["Backend", "FastAPI", "0.136+", "Async web framework"],
            ["Backend", "SQLAlchemy", "2.0", "ORM and database toolkit"],
            ["Backend", "Pydantic", "2.7+", "Data validation"],
            ["Backend", "Uvicorn", "0.27+", "ASGI server"],
            ["AI/ML", "LangGraph", "Latest", "Agent state machine orchestration"],
            ["AI/ML", "LangChain", "Latest", "LLM integration framework"],
            ["AI/ML", "Ollama", "0.1.7+", "Local LLM serving (llama3.2)"],
            ["Database", "PostgreSQL / SQLite", "15+ / 3.x", "Primary data store"],
            ["Cache", "Redis", "5.0+", "Pub/Sub and caching"],
            ["Storage", "MinIO", "7.2+", "S3-compatible object storage"],
            ["Security Tools", "sqlmap", "Latest", "SQL injection testing"],
            ["Security Tools", "Nuclei", "Latest", "Template-based vulnerability scanning"],
            ["Security Tools", "Katana", "Latest", "Web crawling and discovery"],
            ["Security Tools", "httpx", "Latest", "HTTP probing and tech detection"],
            ["Security Tools", "Playwright", "1.41+", "Browser automation and screenshots"],
            ["Monitoring", "Prometheus", "Latest", "Metrics collection"],
            ["Monitoring", "structlog", "24.1+", "Structured logging"],
        ],
        col_widths=[2.5, 4.0, 2.5, 7.0],
    )

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # 18. DEPLOYMENT ARCHITECTURE
    # ════════════════════════════════════════════════════════════════════════
    doc.add_heading("18. Deployment Architecture", level=1)

    doc.add_heading("18.1 Local Development", level=2)
    doc.add_paragraph("For local development, the platform runs with:")
    local_items = [
        "Backend: uvicorn app.main:app --reload --host 0.0.0.0 --port 8000",
        "Frontend: npm run dev (Vite dev server on port 5173)",
        "Database: SQLite (local file, zero configuration)",
        "LLM: Ollama running locally (ollama serve)",
        "All services can be started with a single command: .\\start-local.ps1",
    ]
    for item in local_items:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("18.2 Production (Docker Compose)", level=2)
    doc.add_paragraph("Production deployment uses Docker Compose with:")
    prod_items = [
        "FastAPI backend container with Gunicorn + Uvicorn workers",
        "React frontend served by Nginx",
        "PostgreSQL container for persistent data",
        "Redis container for caching and Pub/Sub",
        "MinIO container for object storage",
        "Ollama container for local LLM inference",
        "Celery workers for background task processing",
    ]
    for item in prod_items:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # 19. GLOSSARY
    # ════════════════════════════════════════════════════════════════════════
    doc.add_heading("19. Glossary of Terms", level=1)
    add_styled_table(doc,
        ["Term", "Definition"],
        [
            ["OWASP", "Open Worldwide Application Security Project — a nonprofit foundation that publishes the industry-standard Top 10 web application security risks"],
            ["IDOR", "Insecure Direct Object Reference — a vulnerability where an attacker can access other users' data by manipulating resource identifiers"],
            ["XSS", "Cross-Site Scripting — injecting malicious scripts into web pages viewed by other users"],
            ["SQLi", "SQL Injection — inserting malicious SQL code into application queries to manipulate the database"],
            ["SSRF", "Server-Side Request Forgery — tricking the server into making requests to internal services"],
            ["CWE", "Common Weakness Enumeration — a categorized list of software and hardware weakness types"],
            ["CVE", "Common Vulnerabilities and Exposures — publicly disclosed cybersecurity vulnerabilities"],
            ["PoC", "Proof of Concept — evidence demonstrating that a vulnerability is real and exploitable"],
            ["LLM", "Large Language Model — AI model used for text generation (e.g., GPT, Llama)"],
            ["LangGraph", "A library for building stateful, multi-agent AI workflows as directed graphs"],
            ["Ollama", "An open-source tool for running LLMs locally on your own hardware"],
            ["False Positive", "A vulnerability reported by a scanner that does not actually exist"],
            ["Zero-Hallucination", "A design principle ensuring AI does not generate fabricated security findings"],
            ["Rule Engine", "A deterministic system that makes decisions based on predefined rules rather than AI inference"],
            ["MinIO", "An S3-compatible object storage system for storing files, screenshots, and reports"],
        ],
        col_widths=[3.0, 13.0],
    )

    # ── Final page ──────────────────────────────────────────────────────────
    doc.add_page_break()
    for _ in range(8):
        doc.add_paragraph()
    end = doc.add_paragraph()
    end.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = end.add_run("— End of Document —")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(150, 150, 150)
    run.italic = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"\nGenerated on {datetime.now().strftime('%B %d, %Y at %I:%M %p')}")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(180, 180, 180)

    return doc


if __name__ == "__main__":
    print("Generating RakshAI Architecture Document...")
    doc = build_document()
    output_path = os.path.join(os.path.dirname(__file__), "RakshAI_Architecture_Client_Guide.docx")
    doc.save(output_path)
    print(f"[OK] Document saved to: {output_path}")
