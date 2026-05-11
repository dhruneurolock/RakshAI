"""
Advanced Report Generation Service

Responsibilities:
1. Generate PDF reports (WeasyPrint)
2. Generate Word documents (python-docx)
3. Generate Excel spreadsheets (openpyxl)
4. LLM-generated executive summary
5. CVSS scoring
6. Compliance mapping (OWASP, ISO 27001, PCI-DSS)
7. Embed screenshots and evidence
8. Upload reports to MinIO
"""

from typing import Dict, Any, List, Optional
import json
import asyncio
from datetime import datetime
from io import BytesIO
import logging

# PDF generation
try:
    from weasyprint import HTML, CSS  # type: ignore
except (ImportError, OSError, Exception):
    HTML = None
    CSS = None

# Word generation
try:
    from docx import Document  # type: ignore
    from docx.shared import Inches, Pt, RGBColor  # type: ignore
    from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore
except ImportError:
    Document = None
    Inches = Pt = RGBColor = None
    WD_ALIGN_PARAGRAPH = None

# Excel generation
try:
    from openpyxl import Workbook  # type: ignore
    from openpyxl.styles import Font, PatternFill, Alignment  # type: ignore
    from openpyxl.utils import get_column_letter  # type: ignore
except ImportError:
    Workbook = None
    Font = PatternFill = Alignment = None
    get_column_letter = None

from app.services.llm_service import get_llm_service
from app.services.storage_service import get_storage_service

logger = logging.getLogger(__name__)


class ReportGeneratorService:
    """Advanced report generation with multiple export formats"""
    
    def __init__(self):
        self.llm_service = None
        self.storage_service = None
    
    async def initialize(self):
        """Initialize services"""
        self.llm_service = await get_llm_service()
        self.storage_service = await get_storage_service()
    
    async def generate_report(
        self,
        scan_id: str,
        scan_data: Dict[str, Any],
        findings: List[Dict[str, Any]],
        format: str = "all"
    ) -> Dict[str, Any]:
        """
        Generate comprehensive penetration testing report
        
        Args:
            scan_id: Scan identifier
            scan_data: Scan metadata
            findings: List of validated findings
            format: Export format (pdf, word, excel, all)
            
        Returns:
            Dict with report URLs
        """
        try:
            logger.info(f"Generating {format} report for scan {scan_id}")

            # Avoid frontend timeout when LLM is slow/unavailable.
            try:
                executive_summary = await asyncio.wait_for(
                    self._generate_executive_summary(scan_data, findings),
                    timeout=8,
                )
            except asyncio.TimeoutError:
                logger.warning("Executive summary generation timed out; using fallback")
                executive_summary = "Executive summary unavailable."

            report_data = {
                "scan_id": scan_id,
                "scan_data": scan_data,
                "findings": findings,
                "executive_summary": executive_summary,
                "generated_at": datetime.utcnow().isoformat(),
                "statistics": self._calculate_statistics(findings),
            }

            urls: Dict[str, str] = {}
            errors: List[str] = []

            if format in ["pdf", "all"]:
                try:
                    pdf_url = await self._generate_pdf(scan_id, report_data)
                    if pdf_url:
                        urls["pdf"] = pdf_url
                    else:
                        errors.append("pdf: empty output")
                except Exception as e:
                    logger.error(f"PDF generation error: {e}")
                    errors.append(f"pdf: {e}")

            if format in ["word", "docx", "all"]:
                try:
                    word_url = await self._generate_word(scan_id, report_data)
                    if word_url:
                        urls["word"] = word_url
                    else:
                        errors.append("word: empty output")
                except Exception as e:
                    logger.error(f"Word generation error: {e}")
                    errors.append(f"word: {e}")

            if format in ["excel", "xlsx", "all"]:
                try:
                    excel_url = await self._generate_excel(scan_id, report_data)
                    if excel_url:
                        urls["excel"] = excel_url
                    else:
                        errors.append("excel: empty output")
                except Exception as e:
                    logger.error(f"Excel generation error: {e}")
                    errors.append(f"excel: {e}")

            if not urls:
                return {
                    "success": False,
                    "error": "; ".join(errors) if errors else "No report output generated",
                }

            return {
                "success": True,
                "scan_id": scan_id,
                "report_urls": urls,
                "generated_at": report_data["generated_at"],
                "errors": errors,
            }

        except Exception as e:
            logger.error(f"Report generation error for scan {scan_id}: {e}")
            return {
                "success": False,
                "error": str(e)
            }
    
    async def _generate_executive_summary(
        self,
        scan_data: Dict[str, Any],
        findings: List[Dict[str, Any]]
    ) -> str:
        """Use LLM to generate executive summary"""
        
        try:
            # Prepare findings summary
            findings_summary = self._summarize_findings(findings)
            
            prompt = f"""Generate an executive summary for this security assessment:

Target: {scan_data.get('target_url')}
Scan Type: {scan_data.get('scan_type')}
Duration: {scan_data.get('duration', 'N/A')}

Findings Summary:
{findings_summary}

Generate a professional executive summary (3-4 paragraphs) suitable for C-level executives covering:
1. Overall security posture
2. Critical findings and business impact
3. Risk level assessment
4. High-level recommendations

Use business language, not technical jargon.
"""
            
            summary = await self.llm_service.analyze(
                prompt=prompt,
                response_format="text",
                use_strategy_model=False
            )
            
            return summary.strip()
            
        except Exception as e:
            logger.error(f"Executive summary generation error: {e}")
            return "Executive summary unavailable."
    
    def _summarize_findings(self, findings: List[Dict[str, Any]]) -> str:
        """Create text summary of findings for LLM"""
        
        severity_counts = {"CRITICAL": 0, "HIGH": 0, "MEDIUM": 0, "LOW": 0, "INFO": 0}
        type_counts = {}
        
        for finding in findings:
            severity = str(finding.get("severity", "LOW")).upper()
            severity_counts[severity] = severity_counts.get(severity, 0) + 1
            
            ftype = finding.get("type")
            type_counts[ftype] = type_counts.get(ftype, 0) + 1
        
        summary_lines = [
            f"Total Findings: {len(findings)}",
            f"Critical: {severity_counts['CRITICAL']}",
            f"High: {severity_counts['HIGH']}",
            f"Medium: {severity_counts['MEDIUM']}",
            f"Low: {severity_counts['LOW']}",
            f"Info: {severity_counts['INFO']}",
            "",
            "Vulnerability Types:"
        ]
        
        for vtype, count in sorted(type_counts.items(), key=lambda x: x[1], reverse=True):
            summary_lines.append(f"  - {vtype}: {count}")
        
        return "\n".join(summary_lines)
    
    def _calculate_statistics(self, findings: List[Dict[str, Any]]) -> Dict[str, Any]:
        """Calculate report statistics"""
        
        total = len(findings)
        severity_counts = {"CRITICAL": 0, "HIGH": 0, "MEDIUM": 0, "LOW": 0, "INFO": 0}
        validated = sum(1 for f in findings if f.get("status") == "VALIDATED")
        
        for finding in findings:
            severity = str(finding.get("severity", "LOW")).upper()
            severity_counts[severity] = severity_counts.get(severity, 0) + 1
        
        return {
            "total_findings": total,
            "validated_findings": validated,
            "severity_breakdown": severity_counts,
            "validation_rate": validated / total if total > 0 else 0
        }
    
    async def _generate_pdf(self, scan_id: str, report_data: Dict[str, Any]) -> str:
        """Generate PDF report using WeasyPrint"""

        try:
            if HTML is None:
                raise RuntimeError("PDF export dependency unavailable (WeasyPrint)")

            # Build HTML content
            html_content = self._build_html_report(report_data)

            # Generate PDF
            pdf_bytes = HTML(string=html_content).write_pdf()
            
            # Upload to MinIO
            url = await self.storage_service.upload_report(
                scan_id,
                pdf_bytes,
                "pdf"
            )
            
            logger.info(f"PDF report generated for scan {scan_id}")
            return url
            
        except Exception as e:
            logger.error(f"PDF generation error: {e}")
            raise
    
    def _build_html_report(self, report_data: Dict[str, Any]) -> str:
        """Build HTML content for PDF report"""
        
        scan_data = report_data["scan_data"]
        findings = report_data["findings"]
        summary = report_data["executive_summary"]
        stats = report_data["statistics"]
        
        html = f"""
<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>RakshAI Security Assessment Report</title>
    <style>
        body {{
            font-family: 'Arial', sans-serif;
            margin: 40px;
            color: #333;
        }}
        h1 {{
            color: #1a73e8;
            border-bottom: 3px solid #1a73e8;
            padding-bottom: 10px;
        }}
        h2 {{
            color: #34a853;
            margin-top: 30px;
        }}
        .severity-critical {{
            color: #d93025;
            font-weight: bold;
        }}
        .severity-high {{
            color: #ea8600;
            font-weight: bold;
        }}
        .severity-medium {{
            color: #f9ab00;
        }}
        .severity-low {{
            color: #1e8e3e;
        }}
        .severity-info {{
            color: #1a73e8;
        }}
        table {{
            width: 100%;
            border-collapse: collapse;
            margin: 20px 0;
        }}
        th, td {{
            border: 1px solid #ddd;
            padding: 12px;
            text-align: left;
        }}
        th {{
            background-color: #1a73e8;
            color: white;
        }}
        .finding-box {{
            border: 1px solid #ddd;
            padding: 15px;
            margin: 20px 0;
            border-radius: 5px;
        }}
        .executive-summary {{
            background-color: #f8f9fa;
            padding: 20px;
            border-left: 4px solid #1a73e8;
            margin: 20px 0;
        }}
    </style>
</head>
<body>
    <h1>Security Assessment Report</h1>
    
    <h2>Scan Information</h2>
    <table>
        <tr>
            <th>Target URL</th>
            <td>{scan_data.get('target_url', 'N/A')}</td>
        </tr>
        <tr>
            <th>Scan Type</th>
            <td>{scan_data.get('scan_type', 'N/A')}</td>
        </tr>
        <tr>
            <th>Generated At</th>
            <td>{report_data.get('generated_at', 'N/A')}</td>
        </tr>
    </table>
    
    <h2>Executive Summary</h2>
    <div class="executive-summary">
        {summary}
    </div>
    
    <h2>Statistics</h2>
    <table>
        <tr>
            <th>Total Findings</th>
            <td>{stats['total_findings']}</td>
        </tr>
        <tr>
            <th>Validated Findings</th>
            <td>{stats['validated_findings']}</td>
        </tr>
        <tr>
            <th>Critical</th>
            <td class="severity-critical">{stats['severity_breakdown']['CRITICAL']}</td>
        </tr>
        <tr>
            <th>High</th>
            <td class="severity-high">{stats['severity_breakdown']['HIGH']}</td>
        </tr>
        <tr>
            <th>Medium</th>
            <td class="severity-medium">{stats['severity_breakdown']['MEDIUM']}</td>
        </tr>
        <tr>
            <th>Low</th>
            <td class="severity-low">{stats['severity_breakdown']['LOW']}</td>
        </tr>
        <tr>
            <th>Info</th>
            <td class="severity-info">{stats['severity_breakdown']['INFO']}</td>
        </tr>
    </table>
    
    <h2>Findings</h2>
"""
        
        # Add each finding
        for i, finding in enumerate(findings, 1):
            severity_class = f"severity-{finding.get('severity', 'low').lower()}"
            
            html += f"""
    <div class="finding-box">
        <h3>#{i} - {finding.get('type', 'Unknown')}</h3>
        <p><strong>Severity:</strong> <span class="{severity_class}">{finding.get('severity', 'N/A')}</span></p>
        <p><strong>URL:</strong> {finding.get('url', 'N/A')}</p>
        <p><strong>Description:</strong> {finding.get('description', 'N/A')}</p>
        <p><strong>Business Impact:</strong> {finding.get('llm_business_impact', 'N/A')[:300]}</p>
        <p><strong>Remediation:</strong> {finding.get('llm_remediation', 'N/A')[:300]}</p>
    </div>
"""
        
        html += """
</body>
</html>
"""
        
        return html
    
    async def _generate_word(self, scan_id: str, report_data: Dict[str, Any]) -> str:
        """Generate Word document report"""

        try:
            if Document is None or WD_ALIGN_PARAGRAPH is None:
                raise RuntimeError("Word export dependency unavailable (python-docx)")

            doc = Document()
            
            # Title
            title = doc.add_heading('Security Assessment Report', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Scan Information
            doc.add_heading('Scan Information', 1)
            scan_table = doc.add_table(rows=3, cols=2)
            scan_table.style = 'Light Grid Accent 1'
            
            scan_data = report_data["scan_data"]
            scan_table.rows[0].cells[0].text = 'Target URL'
            scan_table.rows[0].cells[1].text = scan_data.get('target_url', 'N/A')
            scan_table.rows[1].cells[0].text = 'Scan Type'
            scan_table.rows[1].cells[1].text = scan_data.get('scan_type', 'N/A')
            scan_table.rows[2].cells[0].text = 'Generated At'
            scan_table.rows[2].cells[1].text = report_data.get('generated_at', 'N/A')
            
            # Executive Summary
            doc.add_page_break()
            doc.add_heading('Executive Summary', 1)
            doc.add_paragraph(report_data["executive_summary"])
            
            # Statistics
            doc.add_heading('Statistics', 1)
            stats = report_data["statistics"]
            stats_table = doc.add_table(rows=8, cols=2)
            stats_table.style = 'Light Grid Accent 1'
            
            stats_table.rows[0].cells[0].text = 'Total Findings'
            stats_table.rows[0].cells[1].text = str(stats['total_findings'])
            stats_table.rows[1].cells[0].text = 'Validated Findings'
            stats_table.rows[1].cells[1].text = str(stats['validated_findings'])
            stats_table.rows[2].cells[0].text = 'Critical'
            stats_table.rows[2].cells[1].text = str(stats['severity_breakdown']['CRITICAL'])
            stats_table.rows[3].cells[0].text = 'High'
            stats_table.rows[3].cells[1].text = str(stats['severity_breakdown']['HIGH'])
            stats_table.rows[4].cells[0].text = 'Medium'
            stats_table.rows[4].cells[1].text = str(stats['severity_breakdown']['MEDIUM'])
            stats_table.rows[5].cells[0].text = 'Low'
            stats_table.rows[5].cells[1].text = str(stats['severity_breakdown']['LOW'])
            stats_table.rows[6].cells[0].text = 'Info'
            stats_table.rows[6].cells[1].text = str(stats['severity_breakdown']['INFO'])
            stats_table.rows[7].cells[0].text = 'Validation Rate'
            stats_table.rows[7].cells[1].text = f"{stats['validation_rate']:.1%}"
            
            # Findings
            doc.add_page_break()
            doc.add_heading('Detailed Findings', 1)
            
            for i, finding in enumerate(report_data["findings"], 1):
                doc.add_heading(f"Finding #{i}: {finding.get('type')}", 2)
                doc.add_paragraph(f"Severity: {finding.get('severity')}")
                doc.add_paragraph(f"URL: {finding.get('url')}")
                doc.add_paragraph(f"Description: {finding.get('description')}")
                doc.add_paragraph(f"Business Impact: {finding.get('llm_business_impact', 'N/A')}")
                doc.add_paragraph(f"Remediation: {finding.get('llm_remediation', 'N/A')}")
                doc.add_paragraph("")
            
            # Save to bytes
            word_bytes = BytesIO()
            doc.save(word_bytes)
            word_bytes.seek(0)
            
            # Upload to MinIO
            url = await self.storage_service.upload_report(
                scan_id,
                word_bytes.read(),
                "docx"
            )
            
            logger.info(f"Word report generated for scan {scan_id}")
            return url
            
        except Exception as e:
            logger.error(f"Word generation error: {e}")
            raise
    
    async def _generate_excel(self, scan_id: str, report_data: Dict[str, Any]) -> str:
        """Generate Excel spreadsheet report with professional formatting"""

        try:
            if Workbook is None or PatternFill is None or Font is None or get_column_letter is None:
                raise RuntimeError("Excel export dependency unavailable (openpyxl)")

            from openpyxl.styles import Border, Side, Alignment  # type: ignore

            def _as_text(value: Any, fallback: str = "N/A") -> str:
                if value is None:
                    return fallback
                if isinstance(value, (list, tuple, set)):
                    parts = [str(item).strip() for item in value if str(item).strip()]
                    return "\n".join(parts) if parts else fallback
                text = str(value).strip()
                return text if text else fallback

            def _truncate(value: Any, limit: int = 32000) -> str:
                text = _as_text(value, "")
                return text[:limit]

            def _first_present(finding: Dict[str, Any], keys: list[str], fallback: str = "N/A") -> str:
                for key in keys:
                    value = finding.get(key)
                    if value not in (None, "", [], {}):
                        return _as_text(value, fallback)
                return fallback

            wb = Workbook()
            
            # Define colors for severity levels
            severity_colors = {
                "CRITICAL": "d93025",  # Red
                "HIGH": "ea8600",      # Orange
                "MEDIUM": "f9ab00",    # Amber
                "LOW": "1e8e3e",       # Green
                "INFO": "1a73e8"       # Blue
            }
            
            # Define border style
            thin_border = Border(
                left=Side(style='thin'),
                right=Side(style='thin'),
                top=Side(style='thin'),
                bottom=Side(style='thin')
            )
            
            # ========== SUMMARY SHEET ==========
            ws_summary = wb.active
            ws_summary.title = "Summary"
            ws_summary.column_dimensions['A'].width = 25
            ws_summary.column_dimensions['B'].width = 50
            
            # Title
            ws_summary['A1'] = 'RakshAI Security Assessment Report'
            ws_summary['A1'].font = Font(size=18, bold=True, color="1a73e8")
            ws_summary.merge_cells('A1:B1')
            
            # Scan Info Section
            ws_summary['A3'] = 'Scan Information'
            ws_summary['A3'].font = Font(size=12, bold=True, color="FFFFFF")
            ws_summary['A3'].fill = PatternFill(start_color="1a73e8", end_color="1a73e8", fill_type="solid")
            ws_summary.merge_cells('A3:B3')
            
            row = 4
            scan_data = report_data["scan_data"]
            scan_info = [
                ('Target URL', scan_data.get('target_url', 'N/A')),
                ('Scan Type', scan_data.get('scan_type', 'N/A')),
                ('Duration (seconds)', scan_data.get('duration', 'N/A')),
                ('Generated At', report_data.get('generated_at', 'N/A')),
            ]
            
            for label, value in scan_info:
                ws_summary[f'A{row}'] = label
                ws_summary[f'B{row}'] = str(value)
                ws_summary[f'A{row}'].font = Font(bold=True)
                ws_summary[f'A{row}'].border = thin_border
                ws_summary[f'B{row}'].border = thin_border
                row += 1
            
            # Statistics Section
            row += 1
            ws_summary[f'A{row}'] = 'Severity Breakdown'
            ws_summary[f'A{row}'].font = Font(size=12, bold=True, color="FFFFFF")
            ws_summary[f'A{row}'].fill = PatternFill(start_color="1a73e8", end_color="1a73e8", fill_type="solid")
            ws_summary.merge_cells(f'A{row}:B{row}')
            row += 1
            
            stats = report_data["statistics"]
            severity_data = [
                ('CRITICAL', stats['severity_breakdown']['CRITICAL'], severity_colors['CRITICAL']),
                ('HIGH', stats['severity_breakdown']['HIGH'], severity_colors['HIGH']),
                ('MEDIUM', stats['severity_breakdown']['MEDIUM'], severity_colors['MEDIUM']),
                ('LOW', stats['severity_breakdown']['LOW'], severity_colors['LOW']),
                ('INFO', stats['severity_breakdown']['INFO'], severity_colors['INFO']),
            ]
            
            for severity, count, color in severity_data:
                ws_summary[f'A{row}'] = severity
                ws_summary[f'B{row}'] = int(count)
                ws_summary[f'A{row}'].fill = PatternFill(start_color=color, end_color=color, fill_type="solid")
                ws_summary[f'A{row}'].font = Font(color="FFFFFF", bold=True)
                ws_summary[f'A{row}'].border = thin_border
                ws_summary[f'B{row}'].border = thin_border
                ws_summary[f'B{row}'].alignment = Alignment(horizontal='center')
                row += 1
            
            ws_summary[f'A{row}'] = 'Total Findings'
            ws_summary[f'B{row}'] = stats['total_findings']
            ws_summary[f'A{row}'].font = Font(bold=True)
            ws_summary[f'A{row}'].border = thin_border
            ws_summary[f'B{row}'].border = thin_border
            ws_summary[f'B{row}'].font = Font(bold=True, size=12)
            
            # ========== VULNERABILITY MATRIX SHEET ==========
            ws_findings = wb.create_sheet("Vulnerability Matrix")
            ws_findings.column_dimensions['A'].width = 32
            ws_findings.column_dimensions['B'].width = 24
            ws_findings.column_dimensions['C'].width = 34
            ws_findings.column_dimensions['D'].width = 48
            ws_findings.column_dimensions['E'].width = 22
            ws_findings.column_dimensions['F'].width = 14
            ws_findings.column_dimensions['G'].width = 12
            ws_findings.column_dimensions['H'].width = 42
            ws_findings.column_dimensions['I'].width = 42
            ws_findings.column_dimensions['J'].width = 40
            ws_findings.column_dimensions['K'].width = 18

            # Header row
            headers = [
                'Vulnerability Name',
                'Known Name',
                'How to Find',
                'Description',
                'CVE',
                'OWASP',
                'Severity',
                'Impact',
                'Solution',
                'Good Roads',
                'OWASP 2025',
            ]
            header_row = 1
            for col, header in enumerate(headers, 1):
                cell = ws_findings.cell(row=header_row, column=col)
                cell.value = header
                cell.fill = PatternFill(start_color="1a73e8", end_color="1a73e8", fill_type="solid")
                cell.font = Font(color="FFFFFF", bold=True)
                cell.border = thin_border
                cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
            
            # Data rows
            findings = report_data.get("findings", [])
            for i, finding in enumerate(findings, 1):
                row = i + 1
                severity = finding.get('severity', 'LOW').upper()
                severity_color = severity_colors.get(severity, "1a73e8")

                good_references = finding.get('references') or finding.get('reference_urls') or finding.get('good_reads')
                if not good_references:
                    references = []
                    cwe_value = finding.get('cwe_id')
                    cve_value = finding.get('cve')
                    if cwe_value:
                        references.append(str(cwe_value))
                    if cve_value:
                        references.append(str(cve_value))
                    if finding.get('owasp_category'):
                        references.append(f"OWASP {finding.get('owasp_category')}")
                    good_references = references

                how_to_find_parts = [
                    _first_present(finding, ['how_to_find', 'finding_method', 'attack_vector', 'endpoint_url'], 'N/A'),
                    _first_present(finding, ['response_evidence', 'evidence', 'llm_explanation'], ''),
                    _first_present(finding, ['poc_curl_command', 'request_payload'], ''),
                ]
                how_to_find = "\n\n".join(part for part in how_to_find_parts if part and part != 'N/A') or 'N/A'

                solution_parts = [
                    _first_present(finding, ['llm_remediation', 'remediation'], ''),
                    _first_present(finding, ['llm_business_impact', 'impact'], ''),
                ]
                solution_text = "\n\n".join(part for part in solution_parts if part)

                owasp_2025 = _first_present(finding, ['owasp_2025', 'owasp_top10_2025', 'owasp_category'], 'N/A')

                row_values = [
                    _first_present(finding, ['title', 'vulnerability_name', 'name', 'type'], 'N/A'),
                    _first_present(finding, ['vulnerability_type', 'known_name', 'attack_vector', 'type'], 'N/A'),
                    _truncate(how_to_find),
                    _truncate(finding.get('description', 'N/A')),
                    _first_present(finding, ['cve', 'cve_id'], 'N/A'),
                    _first_present(finding, ['owasp_category'], 'N/A'),
                    severity,
                    _truncate(_first_present(finding, ['llm_business_impact', 'impact'], 'N/A')),
                    _truncate(solution_text or _first_present(finding, ['remediation', 'llm_remediation'], 'N/A')),
                    _truncate(good_references),
                    owasp_2025,
                ]

                for col, value in enumerate(row_values, 1):
                    ws_findings.cell(row=row, column=col, value=value)
                
                for col in range(1, len(headers) + 1):
                    cell = ws_findings.cell(row=row, column=col)
                    cell.border = thin_border
                    if col == 7:  # Severity column
                        cell.fill = PatternFill(start_color=severity_color, end_color=severity_color, fill_type="solid")
                        cell.font = Font(color="FFFFFF", bold=True)
                    cell.alignment = Alignment(vertical='top', wrap_text=True)

            ws_findings.freeze_panes = "A2"
            ws_findings.auto_filter.ref = f"A1:K{max(len(findings) + 1, 1)}"
            
            # ========== REMEDIATION SHEET ==========
            if findings and any(f.get('remediation') for f in findings):
                ws_remediation = wb.create_sheet("Remediation")
                ws_remediation.column_dimensions['A'].width = 5
                ws_remediation.column_dimensions['B'].width = 25
                ws_remediation.column_dimensions['C'].width = 65
                
                headers = ['#', 'Vulnerability', 'Remediation Steps']
                for col, header in enumerate(headers, 1):
                    cell = ws_remediation.cell(row=1, column=col)
                    cell.value = header
                    cell.fill = PatternFill(start_color="1a73e8", end_color="1a73e8", fill_type="solid")
                    cell.font = Font(color="FFFFFF", bold=True)
                    cell.border = thin_border
                
                for i, finding in enumerate(findings, 1):
                    if finding.get('remediation'):
                        row = i + 1
                        ws_remediation.cell(row=row, column=1, value=i)
                        ws_remediation.cell(row=row, column=2, value=finding.get('type', 'N/A'))
                        ws_remediation.cell(row=row, column=3, value=finding.get('remediation', 'N/A')[:1000])
                        
                        for col in range(1, 4):
                            cell = ws_remediation.cell(row=row, column=col)
                            cell.border = thin_border
                            cell.alignment = Alignment(vertical='top', wrap_text=True)
            
            # Save to bytes
            excel_bytes = BytesIO()
            wb.save(excel_bytes)
            excel_bytes.seek(0)
            
            # Upload to MinIO
            url = await self.storage_service.upload_report(
                scan_id,
                excel_bytes.read(),
                "xlsx"
            )
            
            logger.info(f"Excel report generated for scan {scan_id}")
            return url
            
        except Exception as e:
            logger.error(f"Excel generation error: {e}")
            raise


# Singleton instance
_report_generator = None

async def get_report_generator() -> ReportGeneratorService:
    """Get report generator service instance"""
    global _report_generator
    if _report_generator is None:
        _report_generator = ReportGeneratorService()
        await _report_generator.initialize()
    return _report_generator
